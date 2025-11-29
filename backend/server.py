from fastapi import FastAPI, APIRouter, UploadFile, File, Form, HTTPException
from dotenv import load_dotenv
from starlette.middleware.cors import CORSMiddleware
from motor.motor_asyncio import AsyncIOMotorClient
import os
import logging
from pathlib import Path
from pydantic import BaseModel, Field, ConfigDict
from typing import List, Optional, Dict, Any
import uuid
from datetime import datetime, timezone
import chromadb
from chromadb.config import Settings
from emergentintegrations.llm.chat import LlmChat, UserMessage
import PyPDF2
from docx import Document as DocxDocument
import io
import asyncio

ROOT_DIR = Path(__file__).parent
load_dotenv(ROOT_DIR / '.env')

# MongoDB connection
mongo_url = os.environ['MONGO_URL']
client = AsyncIOMotorClient(mongo_url)
db = client[os.environ['DB_NAME']]

# ChromaDB setup
chroma_client = chromadb.Client(Settings(
    persist_directory="/app/backend/chroma_db",
    anonymized_telemetry=False
))

# Create collections
try:
    hr_collection = chroma_client.get_or_create_collection(name="hr_documents")
    resume_collection = chroma_client.get_or_create_collection(name="resumes")
except:
    hr_collection = chroma_client.get_collection(name="hr_documents")
    resume_collection = chroma_client.get_collection(name="resumes")

# Create the main app
app = FastAPI()
api_router = APIRouter(prefix="/api")

# Get API key
API_KEY = os.environ.get('EMERGENT_LLM_KEY')

# Define Models
class Document(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    filename: str
    content: str
    doc_type: str  # 'hr', 'policy', 'general'
    uploaded_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

class ChatMessage(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    session_id: str
    role: str  # 'user' or 'assistant'
    message: str
    confidence: Optional[float] = None
    escalated: bool = False
    timestamp: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

class Resume(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    filename: str
    content: str
    candidate_name: Optional[str] = None
    score: Optional[float] = None
    analysis: Optional[str] = None
    uploaded_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

class InterviewSession(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    candidate_name: str
    position: str
    messages: List[Dict[str, Any]] = []
    evaluation: Optional[str] = None
    score: Optional[float] = None
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

class ChatRequest(BaseModel):
    session_id: str
    message: str

class ResumeScreenRequest(BaseModel):
    job_description: str

class InterviewRequest(BaseModel):
    session_id: str
    candidate_name: str
    position: str
    message: Optional[str] = None

# Helper functions
def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract text from PDF"""
    try:
        pdf_file = io.BytesIO(file_bytes)
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text()
        return text
    except Exception as e:
        logging.error(f"Error extracting PDF text: {e}")
        return ""

def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text from DOCX"""
    try:
        doc = DocxDocument(io.BytesIO(file_bytes))
        text = ""
        for para in doc.paragraphs:
            text += para.text + "\n"
        return text
    except Exception as e:
        logging.error(f"Error extracting DOCX text: {e}")
        return ""

async def query_knowledge_base(query: str, collection, top_k: int = 3) -> List[str]:
    """Query ChromaDB for relevant documents"""
    try:
        results = collection.query(
            query_texts=[query],
            n_results=top_k
        )
        if results and results['documents']:
            return results['documents'][0]
        return []
    except Exception as e:
        logging.error(f"Error querying knowledge base: {e}")
        return []

async def get_ai_response(messages: List[Dict[str, str]], system_message: str) -> tuple[str, float]:
    """Get response from AI with confidence score"""
    try:
        chat = LlmChat(
            api_key=API_KEY,
            session_id=str(uuid.uuid4()),
            system_message=system_message
        ).with_model("openai", "gpt-5.1")
        
        # For simplicity, just send the last user message
        last_message = messages[-1]['content'] if messages else ""
        user_msg = UserMessage(text=last_message)
        
        response = await chat.send_message(user_msg)
        
        # Simple confidence calculation (can be enhanced)
        confidence = 0.85 if len(response) > 50 else 0.60
        
        return response, confidence
    except Exception as e:
        logging.error(f"Error getting AI response: {e}")
        return "I apologize, but I'm having trouble processing your request. Please try again.", 0.5

# API Routes
@api_router.get("/")
async def root():
    return {"message": "AI HR & Support Agent API"}

@api_router.post("/upload-document")
async def upload_document(
    file: UploadFile = File(...),
    doc_type: str = Form("hr")
):
    """Upload HR documents to knowledge base"""
    try:
        file_bytes = await file.read()
        
        # Extract text based on file type
        if file.filename.endswith('.pdf'):
            text = extract_text_from_pdf(file_bytes)
        elif file.filename.endswith('.docx'):
            text = extract_text_from_docx(file_bytes)
        else:
            raise HTTPException(status_code=400, detail="Unsupported file format")
        
        if not text:
            raise HTTPException(status_code=400, detail="Could not extract text from document")
        
        # Store in ChromaDB
        doc_id = str(uuid.uuid4())
        hr_collection.add(
            documents=[text],
            metadatas=[{"filename": file.filename, "doc_type": doc_type}],
            ids=[doc_id]
        )
        
        # Store in MongoDB
        doc = Document(
            id=doc_id,
            filename=file.filename,
            content=text[:500],  # Store preview
            doc_type=doc_type
        )
        doc_dict = doc.model_dump()
        doc_dict['uploaded_at'] = doc_dict['uploaded_at'].isoformat()
        
        await db.documents.insert_one(doc_dict)
        
        return {"message": "Document uploaded successfully", "doc_id": doc_id}
    except Exception as e:
        logging.error(f"Error uploading document: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@api_router.get("/documents")
async def get_documents():
    """Get all uploaded documents"""
    try:
        docs = await db.documents.find({}, {"_id": 0}).to_list(1000)
        for doc in docs:
            if isinstance(doc['uploaded_at'], str):
                doc['uploaded_at'] = datetime.fromisoformat(doc['uploaded_at'])
        return {"documents": docs}
    except Exception as e:
        logging.error(f"Error fetching documents: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@api_router.post("/chat")
async def chat(request: ChatRequest):
    """Chat with HR assistant"""
    try:
        # Store user message
        user_msg = ChatMessage(
            session_id=request.session_id,
            role="user",
            message=request.message
        )
        user_msg_dict = user_msg.model_dump()
        user_msg_dict['timestamp'] = user_msg_dict['timestamp'].isoformat()
        await db.chat_messages.insert_one(user_msg_dict)
        
        # Query knowledge base
        relevant_docs = await query_knowledge_base(request.message, hr_collection)
        
        # Build context
        context = "\n\n".join(relevant_docs) if relevant_docs else "No relevant documents found."
        
        system_message = f"""You are an AI HR & Support Assistant. Your responsibilities:

1. Answer HR-related queries about leaves, attendance, payroll, benefits, policies, onboarding
2. Provide general customer support
3. Be professional, friendly, and supportive
4. Respond concisely (3-6 lines) unless more detail is requested

Knowledge Base Context:
{context}

Rules:
- Use the knowledge base to answer questions
- If information is not in the knowledge base and you're not confident, say: "I will escalate this to the HR team. Can I get your employee ID and contact details?"
- Never invent company-specific information
- Ask clarifying questions if needed"""
        
        # Get AI response
        messages = [{"role": "user", "content": request.message}]
        response_text, confidence = await get_ai_response(messages, system_message)
        
        # Determine if escalation needed
        escalated = confidence < 0.60
        if escalated and "escalate" not in response_text.lower():
            response_text = "I will escalate this request to the HR/Support team. Can I get your employee ID and contact details?"
        
        # Store assistant message
        assistant_msg = ChatMessage(
            session_id=request.session_id,
            role="assistant",
            message=response_text,
            confidence=confidence,
            escalated=escalated
        )
        assistant_msg_dict = assistant_msg.model_dump()
        assistant_msg_dict['timestamp'] = assistant_msg_dict['timestamp'].isoformat()
        await db.chat_messages.insert_one(assistant_msg_dict)
        
        return {
            "response": response_text,
            "confidence": confidence,
            "escalated": escalated
        }
    except Exception as e:
        logging.error(f"Error in chat: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@api_router.get("/chat-history/{session_id}")
async def get_chat_history(session_id: str):
    """Get chat history for a session"""
    try:
        messages = await db.chat_messages.find(
            {"session_id": session_id},
            {"_id": 0}
        ).sort("timestamp", 1).to_list(1000)
        
        for msg in messages:
            if isinstance(msg['timestamp'], str):
                msg['timestamp'] = datetime.fromisoformat(msg['timestamp'])
        
        return {"messages": messages}
    except Exception as e:
        logging.error(f"Error fetching chat history: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@api_router.post("/upload-resume")
async def upload_resume(file: UploadFile = File(...)):
    """Upload resume for screening"""
    try:
        file_bytes = await file.read()
        
        # Extract text
        if file.filename.endswith('.pdf'):
            text = extract_text_from_pdf(file_bytes)
        elif file.filename.endswith('.docx'):
            text = extract_text_from_docx(file_bytes)
        else:
            raise HTTPException(status_code=400, detail="Unsupported file format")
        
        if not text:
            raise HTTPException(status_code=400, detail="Could not extract text from resume")
        
        # Store in ChromaDB
        resume_id = str(uuid.uuid4())
        resume_collection.add(
            documents=[text],
            metadatas=[{"filename": file.filename}],
            ids=[resume_id]
        )
        
        # Store in MongoDB
        resume = Resume(
            id=resume_id,
            filename=file.filename,
            content=text[:1000]
        )
        resume_dict = resume.model_dump()
        resume_dict['uploaded_at'] = resume_dict['uploaded_at'].isoformat()
        
        await db.resumes.insert_one(resume_dict)
        
        return {"message": "Resume uploaded successfully", "resume_id": resume_id}
    except Exception as e:
        logging.error(f"Error uploading resume: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@api_router.post("/screen-resumes")
async def screen_resumes(request: ResumeScreenRequest):
    """Screen all resumes against job description"""
    try:
        # Get all resumes
        resumes = await db.resumes.find({}, {"_id": 0}).to_list(1000)
        
        if not resumes:
            return {"message": "No resumes to screen", "results": []}
        
        results = []
        
        for resume in resumes:
            system_message = f"""You are a Resume Screening AI. Analyze the resume against the job description and provide:

1. Match score (0-100)
2. Key strengths
3. Gaps or concerns
4. Recommendation (Strong Match / Good Match / Weak Match / No Match)

Job Description:
{request.job_description}

Resume:
{resume['content']}

Provide a concise analysis."""
            
            messages = [{"role": "user", "content": "Please analyze this resume."}]
            analysis, _ = await get_ai_response(messages, system_message)
            
            # Extract score (simple heuristic)
            score = 75.0  # Default
            if "Strong Match" in analysis:
                score = 90.0
            elif "Good Match" in analysis:
                score = 75.0
            elif "Weak Match" in analysis:
                score = 50.0
            elif "No Match" in analysis:
                score = 30.0
            
            # Update resume in DB
            await db.resumes.update_one(
                {"id": resume['id']},
                {"$set": {"score": score, "analysis": analysis}}
            )
            
            results.append({
                "resume_id": resume['id'],
                "filename": resume['filename'],
                "score": score,
                "analysis": analysis
            })
            
            # Add small delay to avoid rate limiting
            await asyncio.sleep(0.5)
        
        # Sort by score
        results.sort(key=lambda x: x['score'], reverse=True)
        
        return {"results": results}
    except Exception as e:
        logging.error(f"Error screening resumes: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@api_router.get("/resumes")
async def get_resumes():
    """Get all resumes"""
    try:
        resumes = await db.resumes.find({}, {"_id": 0}).to_list(1000)
        for resume in resumes:
            if isinstance(resume['uploaded_at'], str):
                resume['uploaded_at'] = datetime.fromisoformat(resume['uploaded_at'])
        return {"resumes": resumes}
    except Exception as e:
        logging.error(f"Error fetching resumes: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@api_router.post("/interview")
async def conduct_interview(request: InterviewRequest):
    """Conduct AI interview"""
    try:
        # Get or create interview session
        session = await db.interview_sessions.find_one({"id": request.session_id}, {"_id": 0})
        
        if not session:
            # Create new session
            new_session = InterviewSession(
                id=request.session_id,
                candidate_name=request.candidate_name,
                position=request.position,
                messages=[]
            )
            session_dict = new_session.model_dump()
            session_dict['created_at'] = session_dict['created_at'].isoformat()
            await db.interview_sessions.insert_one(session_dict)
            session = session_dict
            
            # Start with introduction
            intro = f"Hello {request.candidate_name}! I'm your AI interviewer today for the {request.position} position. Let's begin with a simple question: Can you tell me about your background and why you're interested in this role?"
            
            session['messages'].append({
                "role": "assistant",
                "content": intro,
                "timestamp": datetime.now(timezone.utc).isoformat()
            })
            
            await db.interview_sessions.update_one(
                {"id": request.session_id},
                {"$set": {"messages": session['messages']}}
            )
            
            return {"response": intro, "question_number": 1}
        
        # Add user response
        if request.message:
            session['messages'].append({
                "role": "user",
                "content": request.message,
                "timestamp": datetime.now(timezone.utc).isoformat()
            })
        
        # Generate next question or evaluation
        question_count = sum(1 for m in session['messages'] if m['role'] == 'assistant')
        
        if question_count >= 5:  # After 5 questions, provide evaluation
            system_message = f"""You are an AI Interview Evaluator. Review the interview conversation and provide:

1. Overall assessment
2. Strengths demonstrated
3. Areas of concern
4. Score (0-100)
5. Hiring recommendation

Position: {request.position}
Candidate: {request.candidate_name}

Interview Transcript:
"""
            for msg in session['messages']:
                system_message += f"\n{msg['role']}: {msg['content']}"
            
            messages = [{"role": "user", "content": "Please evaluate this interview."}]
            evaluation, _ = await get_ai_response(messages, system_message)
            
            # Extract score
            score = 75.0  # Default
            
            await db.interview_sessions.update_one(
                {"id": request.session_id},
                {"$set": {
                    "messages": session['messages'],
                    "evaluation": evaluation,
                    "score": score
                }}
            )
            
            return {
                "response": evaluation,
                "completed": True,
                "score": score
            }
        
        # Generate next question
        system_message = f"""You are an AI Interviewer for a {request.position} position. Based on the conversation so far, ask the next relevant interview question. 

Make questions:
- Relevant to the position
- Progressive (build on previous answers)
- Mix of technical and behavioral
- Clear and professional

Previous conversation:
"""
        for msg in session['messages'][-4:]:  # Last 4 messages for context
            system_message += f"\n{msg['role']}: {msg['content']}"
        
        messages = [{"role": "user", "content": "Ask the next interview question."}]
        next_question, _ = await get_ai_response(messages, system_message)
        
        session['messages'].append({
            "role": "assistant",
            "content": next_question,
            "timestamp": datetime.now(timezone.utc).isoformat()
        })
        
        await db.interview_sessions.update_one(
            {"id": request.session_id},
            {"$set": {"messages": session['messages']}}
        )
        
        return {
            "response": next_question,
            "question_number": question_count + 1,
            "completed": False
        }
    except Exception as e:
        logging.error(f"Error in interview: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@api_router.get("/interview/{session_id}")
async def get_interview_session(session_id: str):
    """Get interview session"""
    try:
        session = await db.interview_sessions.find_one({"id": session_id}, {"_id": 0})
        if not session:
            raise HTTPException(status_code=404, detail="Interview session not found")
        
        if isinstance(session['created_at'], str):
            session['created_at'] = datetime.fromisoformat(session['created_at'])
        
        return session
    except HTTPException:
        raise
    except Exception as e:
        logging.error(f"Error fetching interview session: {e}")
        raise HTTPException(status_code=500, detail=str(e))

# Include the router in the main app
app.include_router(api_router)

app.add_middleware(
    CORSMiddleware,
    allow_credentials=True,
    allow_origins=os.environ.get('CORS_ORIGINS', '*').split(','),
    allow_methods=["*"],
    allow_headers=["*"],
)

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

@app.on_event("shutdown")
async def shutdown_db_client():
    client.close()